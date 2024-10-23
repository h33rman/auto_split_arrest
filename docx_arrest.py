
import os
import re
from docx import Document


# Function to save a document into a new file
def save_new_doc(new_arrest, arrest_number):
    output_dir = "output"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Create the output file name based on the arrest number
    new_filename = f"{output_dir}/arrest_number_{arrest_number}.docx"

    # Create a new document and add the paragraphs
    new_doc = Document()
    for para in new_arrest:
        new_doc.add_paragraph(para.text)
    new_doc.save(new_filename)

    print(f"Saved: {new_filename}")


# Function to extract the arrest number from the paragraph that starts with "Arrêt n°"
def extract_arrest_number(para_text):
    # Use regular expression to find the arrest number in the format "Arrêt n°123"
    match = re.search(r'Arrêt n°\s*(\d+)', para_text)
    if match:
        return match.group(1)  # Return the number part
    return None


# Function to split the document based on "Arrêt n°"
def split_document(doc_path):
    doc = Document(doc_path)
    new_arrest = []
    arrest_number = None  # Keep track of the current arrest number

    for para in doc.paragraphs:
        # Detect the beginning of a new "Arrêt" using "Arrêt n°"
        if para.text.startswith("Arrêt n°"):
            # If we have an arrest collected, save it to a file
            if new_arrest and arrest_number:
                save_new_doc(new_arrest, arrest_number)
                new_arrest = []

            # Extract the arrest number from the current paragraph
            arrest_number = extract_arrest_number(para.text)

        # Add a paragraph to the current arrest content
        new_arrest.append(para)

    # Save the last "Arrêt" if present
    if new_arrest and arrest_number:
        save_new_doc(new_arrest, arrest_number)


# Main script to run the split process
if __name__ == "__main__":
    input_doc = "/home/herman/Downloads/Documents/LivreRecueil_2009_all.docx"
    split_document(input_doc)
